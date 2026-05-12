"""
converter.py — Conversion PDF → DOCX
Utilise pdf2docx (wrapper LibreOffice-quality).
Fallback automatique via pdfminer + python-docx si pdf2docx absent.
"""

from pathlib import Path
import traceback


# ──────────────────────────────────────────
# Primary engine: pdf2docx
# ──────────────────────────────────────────
def _convert_with_pdf2docx(pdf_path: Path, output_path: Path) -> tuple[bool, str | None]:
    try:
        from pdf2docx import Converter
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path), start=0, end=None)
        cv.close()
        return True, None
    except ImportError:
        return False, "pdf2docx not installed"
    except Exception as e:
        return False, traceback.format_exc()


# ──────────────────────────────────────────
# Fallback engine: pdfminer + python-docx
# ──────────────────────────────────────────
def _convert_with_fallback(pdf_path: Path, output_path: Path) -> tuple[bool, str | None]:
    """
    Fallback minimal : extrait le texte brut du PDF et le place dans un .docx.
    Pas de mise en forme avancée — utilisé seulement si pdf2docx est absent.
    """
    try:
        from pdfminer.high_level import extract_text
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        text = extract_text(str(pdf_path))
        if not text.strip():
            return False, "Could not extract any text from this PDF (scanned image?)."

        doc = Document()

        # Title placeholder
        title = doc.add_heading("Converted Document", level=1)
        title.runs[0].font.color.rgb = RGBColor(0x4F, 0x6E, 0xF7)

        # Metadata paragraph
        info = doc.add_paragraph()
        info.add_run(f"Source file: {pdf_path.name}").italic = True
        info.runs[0].font.color.rgb = RGBColor(0x7A, 0x7F, 0x9A)
        doc.add_paragraph()   # blank line

        # Content — split by line, detect headings heuristically
        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            # Heuristic: short ALL-CAPS lines → heading
            if stripped.isupper() and len(stripped) < 80:
                p = doc.add_heading(stripped, level=2)
            else:
                p = doc.add_paragraph(stripped)
                p.runs[0].font.size = Pt(11) if p.runs else None

        doc.save(str(output_path))
        return True, None

    except ImportError as e:
        return False, f"Missing dependency: {e}. Install: pip install pdfminer.six python-docx"
    except Exception:
        return False, traceback.format_exc()


# ──────────────────────────────────────────
# Public entry point
# ──────────────────────────────────────────
def convert_pdf_to_docx(pdf_path: Path, output_path: Path) -> tuple[bool, str | None]:
    """
    Convert *pdf_path* to *output_path* (.docx).
    Returns (success: bool, error_message: str | None).
    """
    if not pdf_path.exists():
        return False, f"Input file not found: {pdf_path}"

    # Try primary engine first
    ok, err = _convert_with_pdf2docx(pdf_path, output_path)
    if ok:
        return True, None

    # If pdf2docx is missing, use fallback; otherwise surface the real error
    if err and "not installed" in err:
        print("[converter] pdf2docx not found — using pdfminer fallback.")
        return _convert_with_fallback(pdf_path, output_path)

    # pdf2docx failed for another reason — still try fallback
    print(f"[converter] pdf2docx failed: {err}\nTrying fallback…")
    ok2, err2 = _convert_with_fallback(pdf_path, output_path)
    if ok2:
        return True, None

    return False, f"Primary error: {err}\n\nFallback error: {err2}"